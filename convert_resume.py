"""
Script to convert resume website (index.html) to PDF, DOCX, and DOC formats.

Dependencies:
- reportlab: pip install reportlab
- python-docx: pip install python-docx
- beautifulsoup4: pip install beautifulsoup4

Usage:
python convert_resume.py

This script reads index.html, converts it to PDF and DOCX, and creates a DOC file as a copy of DOCX.
"""

import os
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from bs4 import BeautifulSoup
import shutil

def html_to_pdf(html_path, output_pdf):
    try:
        with open(html_path, 'r', encoding='utf-8') as f:
            html_content = f.read()

        soup = BeautifulSoup(html_content, 'html.parser')

        doc = SimpleDocTemplate(output_pdf, pagesize=letter)
        styles = getSampleStyleSheet()

        # Custom styles with colors
        title_style = ParagraphStyle('Title', parent=styles['Title'], fontSize=24, textColor=colors.HexColor('#22d3ee'), alignment=1)
        heading2_style = ParagraphStyle('Heading2', parent=styles['Heading2'], textColor=colors.HexColor('#22d3ee'))
        normal_style = styles['Normal']

        story = []

        # Add profile picture
        img_path = 'WhatsApp Image 2025-09-12 at 19.42.07_7949928b.jpg'
        if os.path.exists(img_path):
            img = Image(img_path, width=2*inch, height=2*inch)
            img.hAlign = 'CENTER'
            story.append(img)
            story.append(Spacer(1, 12))

        # Add name and title from hero
        hero = soup.find(id='hero')
        if hero:
            h1 = hero.find('h1')
            if h1:
                story.append(Paragraph(h1.get_text(strip=True), title_style))
                story.append(Spacer(1, 12))
            h2 = hero.find('h2')
            if h2:
                story.append(Paragraph(h2.get_text(strip=True), heading2_style))
                story.append(Spacer(1, 12))
            tagline = hero.find(class_='tagline')
            if tagline:
                story.append(Paragraph(tagline.get_text(strip=True), normal_style))
                story.append(Spacer(1, 12))
            desc = hero.find(class_='hero-description')
            if desc:
                story.append(Paragraph(desc.get_text(strip=True), normal_style))
                story.append(Spacer(1, 12))

        # Add sections
        for section in soup.find_all('section'):
            h2 = section.find('h2')
            if h2:
                story.append(Paragraph(h2.get_text(strip=True), heading2_style))
                story.append(Spacer(1, 12))
            # Add content
            for p in section.find_all('p'):
                story.append(Paragraph(p.get_text(strip=True), normal_style))
                story.append(Spacer(1, 6))
            for ul in section.find_all('ul'):
                for li in ul.find_all('li'):
                    story.append(Paragraph(f"• {li.get_text(strip=True)}", normal_style))
                    story.append(Spacer(1, 6))

        doc.build(story)
        print(f"PDF generated: {output_pdf}")
    except Exception as e:
        print(f"Error generating PDF: {e}")

def html_to_docx(html_path, output_docx):
    try:
        with open(html_path, 'r', encoding='utf-8') as f:
            html_content = f.read()

        soup = BeautifulSoup(html_content, 'html.parser')

        doc = Document()

        # Add profile picture
        img_path = 'WhatsApp Image 2025-09-12 at 19.42.07_7949928b.jpg'
        if os.path.exists(img_path):
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run()
            run.add_picture(img_path, width=Inches(2))

        # Add name and title from hero
        hero = soup.find(id='hero')
        if hero:
            h1 = hero.find('h1')
            if h1:
                p = doc.add_heading(h1.get_text(strip=True), level=0)
                for run in p.runs:
                    run.font.color.rgb = RGBColor(34, 211, 238)  # #22d3ee
            h2 = hero.find('h2')
            if h2:
                p = doc.add_heading(h2.get_text(strip=True), level=1)
                for run in p.runs:
                    run.font.color.rgb = RGBColor(34, 211, 238)
            tagline = hero.find(class_='tagline')
            if tagline:
                doc.add_paragraph(tagline.get_text(strip=True))
            desc = hero.find(class_='hero-description')
            if desc:
                doc.add_paragraph(desc.get_text(strip=True))

        # Add sections
        for section in soup.find_all('section'):
            h2 = section.find('h2')
            if h2:
                p = doc.add_heading(h2.get_text(strip=True), level=1)
                for run in p.runs:
                    run.font.color.rgb = RGBColor(34, 211, 238)
            # Add content
            for p in section.find_all('p'):
                doc.add_paragraph(p.get_text(strip=True))
            for ul in section.find_all('ul'):
                for li in ul.find_all('li'):
                    doc.add_paragraph(f"• {li.get_text(strip=True)}", style='List Bullet')

        doc.save(output_docx)
        print(f"DOCX generated: {output_docx}")
    except Exception as e:
        print(f"Error generating DOCX: {e}")

def create_doc_from_docx(docx_path, doc_path):
    try:
        # Simple approach: copy docx file with .doc extension
        shutil.copyfile(docx_path, doc_path)
        print(f"DOC generated (copy of DOCX): {doc_path}")
    except Exception as e:
        print(f"Error generating DOC: {e}")

def main():
    html_path = 'index.html'
    output_pdf = 'resume.pdf'
    output_docx = 'resume.docx'
    output_doc = 'resume.doc'

    if not os.path.exists(html_path):
        print(f"Error: {html_path} not found.")
        return

    html_to_pdf(html_path, output_pdf)
    html_to_docx(html_path, output_docx)
    create_doc_from_docx(output_docx, output_doc)

if __name__ == '__main__':
    main()
